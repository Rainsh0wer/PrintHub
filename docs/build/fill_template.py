# -*- coding: utf-8 -*-
"""
Fill the reference PRN232 template with PrintHub content WITHOUT regenerating
styles.

Technique ("template-as-container"): open the real template .docx with
python-docx. That preserves the template's styles.xml, theme, cover page, TOC
field, headers/footers, page setup and table styling automatically. We then:
  1. swap the cover title + author text (in place, keeping run formatting),
  2. delete only the old body (from the first "1." Heading1 to the final
     section properties), keeping the cover + Table of Contents,
  3. rebuild the PrintHub body using the template's own Heading1/Heading2 styles
     and bordered tables.

Result: PrintHub content in the template's exact look. Open in Word and press
Ctrl+A then F9 to refresh the Table of Contents page numbers (the section titles
are unchanged, so the entries stay correct).

Run:  py docs/build/fill_template.py
"""
import os
import re
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(os.path.dirname(HERE))
TEMPLATE = os.path.join(ROOT, "PRN232 - SE1909 .NET _ Final Project Document _ HE190154 _ Phan Mạnh Hùng.docx")
OUTPUT = os.path.join(ROOT, "docs", "docx", "PrintHub_Final_Project_Document.docx")

COVER_TITLE = "PrintHub — Multi-Vendor Printing & Fabrication Service Marketplace"
COVER_AUTHOR = "Ngô Minh Nhật — PRN232 .NET Final Project"
OLD_TITLE = "Student Event Registration System"
OLD_AUTHOR_MARK = "Phan Mạnh Hùng"


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #
def para_text(p_el):
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def replace_run_text(p, new_text):
    """Replace a paragraph's text but keep its paragraph style + first run font."""
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def style_by_id(doc, style_id):
    for s in doc.styles:
        try:
            if s.style_id == style_id:
                return s
        except Exception:
            pass
    return None


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


def shade_cell(cell, fill="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# --------------------------------------------------------------------------- #
# Body building
# --------------------------------------------------------------------------- #
class Builder:
    def __init__(self, doc):
        self.doc = doc
        self.h1 = style_by_id(doc, "Heading1")
        self.h2 = style_by_id(doc, "Heading2")

    def heading(self, text, level=1):
        p = self.doc.add_paragraph()
        style = self.h1 if level == 1 else self.h2
        if style is not None:
            p.style = style
        else:
            p.style = self.doc.styles["Heading %d" % level]
        p.add_run(text)
        return p

    def para(self, text):
        return self.doc.add_paragraph(text)

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(it)
            p.paragraph_format.left_indent = Pt(18)
            # a simple bullet prefix keeps it style-independent
            if p.runs:
                p.runs[0].text = "•  " + p.runs[0].text
        return

    def mono(self, text):
        """A single-cell bordered box for code/request samples (matches template)."""
        t = self.doc.add_table(rows=1, cols=1)
        set_table_borders(t)
        cell = t.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        return t

    def table(self, rows, header=True):
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        set_table_borders(t)
        t.autofit = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                p = cell.paragraphs[0]
                run = p.add_run("" if val is None else str(val))
                if header and ri == 0:
                    run.bold = True
                    shade_cell(cell)
        return t


# --------------------------------------------------------------------------- #
# Cover + body clearing
# --------------------------------------------------------------------------- #
def swap_cover(doc):
    for p in doc.paragraphs:
        txt = para_text(p._p)
        if OLD_TITLE in txt:
            replace_run_text(p, COVER_TITLE)
        elif OLD_AUTHOR_MARK in txt:
            replace_run_text(p, COVER_AUTHOR)


def clear_body_keep_cover_and_toc(doc):
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))

    start = None
    for child in list(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is None:
                continue
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) == "Heading1":
                if para_text(child).strip().startswith("1."):
                    start = child
                    break
    if start is None:
        raise RuntimeError("Could not locate the first '1.' Heading1 to clear from.")

    removing = False
    to_remove = []
    for child in list(body):
        if child is start:
            removing = True
        if removing:
            if child is sectPr:
                break
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)


# --------------------------------------------------------------------------- #
# main
# --------------------------------------------------------------------------- #
def sanitize_docx(path):
    """The template stores some section measurements as floats
    (e.g. '1440.0000000000002'), which python-docx cannot parse as twips.
    Truncate float measurement attributes to integers, preserving all parts."""
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        parts = {n: z.read(n) for n in names}

    xml = parts["word/document.xml"].decode("utf-8")
    xml = re.sub(
        r'(w:(?:top|left|right|bottom|header|footer|gutter|w|h|start|end)=")(\d+)\.\d+(")',
        r"\1\2\3", xml)
    parts["word/document.xml"] = xml.encode("utf-8")

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, parts[n])


def main():
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    sanitize_docx(OUTPUT)

    doc = Document(OUTPUT)
    swap_cover(doc)
    clear_body_keep_cover_and_toc(doc)

    b = Builder(doc)
    from printhub_content import build_content
    build_content(b)

    doc.save(OUTPUT)
    print("Wrote:", OUTPUT)
    print("Tables:", len(doc.tables), "| Paragraphs:", len(doc.paragraphs))
    print("Open in Word, Ctrl+A then F9 to refresh the Table of Contents.")


if __name__ == "__main__":
    import sys
    sys.path.insert(0, HERE)
    main()
