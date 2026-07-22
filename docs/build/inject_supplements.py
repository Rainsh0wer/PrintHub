# -*- coding: utf-8 -*-
"""
Inject the missing pieces into the user's Software Specification docx WITHOUT
regenerating styles (same python-docx "template-as-container" technique):

  * fix: System Context legend "Circle" -> "Rectangle"
  * fix: Use Case List header (Actor/Use case/Expected result -> Use Case/Feature/Actor/Description)
  * add §4.4 Data Dictionary (all fields per table, parsed from docs/diagrams/erd.mmd)
  * add §6.3 Complaint Workflow (state table)
  * replace §9 API Endpoint List with the complete list (parsed from docs/api-endpoints-full.md)

Original file is left untouched; output is a new *_v2.docx.
Run:  py docs/build/inject_supplements.py
"""
import os
import re
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(os.path.dirname(HERE))
SRC = os.path.join(ROOT, "PRN232_NhatNM_Software Specification.docx")
OUT = os.path.join(ROOT, "docs", "docx", "PRN232_NhatNM_Software Specification_v2.docx")
ERD = os.path.join(ROOT, "docs", "diagrams", "erd.mmd")
API = os.path.join(ROOT, "docs", "api-endpoints-full.md")


# ----------------------------------------------------------------------- helpers
def sanitize_docx(path):
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        parts = {n: z.read(n) for n in names}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = re.sub(r'(w:(?:top|left|right|bottom|header|footer|gutter|w|h|start|end)=")(\d+)\.\d+(")', r"\1\2\3", xml)
    parts["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, parts[n])


def style_by_id(doc, style_id):
    for s in doc.styles:
        try:
            if s.style_id == style_id:
                return s
        except Exception:
            pass
    return None


def set_borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
        b.append(e)
    tblPr.append(b)


def shade(cell, fill="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell(cell, text, bold=False, size=None):
    # clear then set
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def para_text(p_el):
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def find_heading(doc, prefix):
    for p in doc.paragraphs:
        if para_text(p._p).strip().startswith(prefix):
            return p
    return None


class Inserter:
    """Inserts headings/paragraphs/tables immediately before an anchor paragraph,
    accumulating in call order, so new content lands at a precise spot in-place."""
    def __init__(self, doc, anchor_para):
        self.doc = doc
        self.anchor = anchor_para
        self.styles = {1: style_by_id(doc, "Heading1"),
                       2: style_by_id(doc, "Heading2"),
                       3: style_by_id(doc, "Heading3")}

    def heading(self, text, level=2):
        style = self.styles.get(level)
        p = self.anchor.insert_paragraph_before("", style=style)
        if style is None:
            p.style = self.doc.styles["Heading %d" % level]
        p.add_run(text)

    def para(self, text):
        self.anchor.insert_paragraph_before(text)

    def bold_para(self, text):
        p = self.anchor.insert_paragraph_before("")
        p.add_run(text).bold = True

    def table(self, rows, header=True, sizes=None):
        ncols = len(rows[0])
        t = self.doc.add_table(rows=len(rows), cols=ncols)
        set_borders(t)
        for ri, row in enumerate(rows):
            # normalise ragged rows to the header width
            if len(row) < ncols:
                row = list(row) + [""] * (ncols - len(row))
            elif len(row) > ncols:
                row = list(row[:ncols - 1]) + [" ".join(str(x) for x in row[ncols - 1:])]
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                set_cell(cell, "" if val is None else str(val), bold=(header and ri == 0),
                         size=(sizes[ci] if sizes else None))
                if header and ri == 0:
                    shade(cell)
        self.anchor._p.addprevious(t._tbl)
        return t


def remove_between(doc, start_para, end_para):
    body = doc.element.body
    removing = False
    to_remove = []
    for child in list(body):
        if child is start_para._p:
            removing = True
            continue
        if child is end_para._p:
            break
        if removing:
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)


# ----------------------------------------------------------------------- content
def parse_erd(path):
    """erd.mmd -> {Table: [(type, field, keymarks), ...]}"""
    tables = {}
    cur = None
    with open(path, encoding="utf-8") as f:
        for line in f:
            s = line.strip()
            m = re.match(r"^(\w+)\s*\{$", s)
            if m:
                cur = m.group(1); tables[cur] = []; continue
            if s == "}":
                cur = None; continue
            if cur:
                parts = s.split()
                if len(parts) >= 2:
                    typ, name = parts[0], parts[1]
                    keys = " ".join(parts[2:]) if len(parts) > 2 else ""
                    tables[cur].append((typ, name, keys))
    return tables


def parse_api_md(path):
    """api md -> [(subsection_title, [rows...]), ...]; skips the '## 9.x' numbering headers."""
    sections = []
    title = None
    rows = []
    with open(path, encoding="utf-8") as f:
        for line in f:
            if line.startswith("## "):
                if title and rows:
                    sections.append((title, rows)); rows = []
                title = line[3:].strip()
            elif line.strip().startswith("|"):
                raw = line.strip().replace("\\|", "")  # protect escaped pipes
                cells = [c.strip().replace("", "|").replace("`", "")
                         for c in raw.strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):  # separator row
                    continue
                rows.append(cells)
        if title and rows:
            sections.append((title, rows))
    return sections


DATA_DICT_PURPOSE = {
    "Users": "Platform accounts (role, wallet, profile).",
    "RefreshTokens": "Server-side refresh tokens for rotation/revocation.",
    "Shops": "Shop storefronts and onboarding status.",
    "ShopStaff": "User×Shop membership (scoped authorization).",
    "ServiceTypes": "Platform service catalogue.",
    "ShopServices": "Shop rate card (Shop×ServiceType with pricing).",
    "PriceRules": "Pricing rules on a rate card entry.",
    "Machines": "Shop production machines.",
    "Materials": "Shop material stock.",
    "DocumentFiles": "Customer document library.",
    "Quotes": "Computed quotes per shop (with expiry).",
    "Orders": "The central transaction aggregate.",
    "OrderItems": "Configured order lines.",
    "OrderStatusHistories": "Append-only order transition log.",
    "WalletTransactions": "Wallet ledger.",
    "Vouchers": "Promotional vouchers.",
    "Reviews": "Shop reviews (one per order).",
    "Complaints": "Dispute/resolution records.",
    "Favourites": "Saved shops.",
    "Notifications": "In-app notifications.",
    "AuditLogs": "Administrative audit trail.",
}


def main():
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    shutil.copyfile(SRC, OUT)
    sanitize_docx(OUT)
    doc = Document(OUT)

    # --- fix 1: legend Circle -> Rectangle (System / Product row only) ---
    for t in doc.tables:
        for row in t.rows:
            texts = [c.text.strip() for c in row.cells]
            if "System / Product" in texts:
                for c in row.cells:
                    if c.text.strip() == "Circle":
                        set_cell(c, "Rectangle")

    # --- fix 2: Use Case List header ---
    for t in doc.tables:
        head = [c.text.strip() for c in t.rows[0].cells]
        if "Expected result" in head and "Descriptions" in head and len(head) == 5:
            new = ["ID", "Use Case", "Feature", "Actor", "Description"]
            for c, v in zip(t.rows[0].cells, new):
                set_cell(c, v, bold=True)

    # --- add §4.4 Data Dictionary (before "5. Business Rules") ---
    anchor5 = find_heading(doc, "5. Business Rules")
    if anchor5 is not None:
        ins = Inserter(doc, anchor5)
        ins.heading("4.4. Data Dictionary", 3)
        ins.para("Every field of each table, generated from the live database schema "
                 "(types shown as SQL types; PK/FK/UK marked). Common audit columns "
                 "CreatedAt/UpdatedAt/CreatedBy/UpdatedBy appear on entities derived "
                 "from the auditable base.")
        erd = parse_erd(ERD)
        for tbl in sorted(erd.keys()):
            ins.bold_para(f"{tbl} — {DATA_DICT_PURPOSE.get(tbl, '')}")
            rows = [["Field", "Type", "Key"]]
            for typ, name, keys in erd[tbl]:
                rows.append([name, typ, keys])
            ins.table(rows, sizes=[9, 9, 9])

    # --- add §6.3 Complaint Workflow (before "7. System Architecture") ---
    anchor7 = find_heading(doc, "7. System Architecture")
    if anchor7 is not None:
        ins = Inserter(doc, anchor7)
        ins.heading("6.3. Complaint / Resolution Workflow", 3)
        ins.table([
            ["Status", "Meaning", "Valid next status"],
            ["Open", "Customer raised a complaint on a completed order", "ShopResponded, Escalated"],
            ["ShopResponded", "Shop proposed a reprint or refund", "Resolved, Escalated"],
            ["Escalated", "Customer rejected, or the shop did not respond in time", "Closed"],
            ["Resolved", "Customer accepted the resolution", "Terminal"],
            ["Closed", "Administrator adjudicated (final ruling)", "Terminal"],
        ])
        ins.para("A complaint escalates automatically when the shop does not respond "
                 "within the window; the administrator's ruling is final within the system.")

    # --- replace §9 with the complete API list ---
    api9 = find_heading(doc, "9. API Endpoint List")
    api10 = find_heading(doc, "10. Security Matrix")
    if api9 is not None and api10 is not None:
        remove_between(doc, api9, api10)
        ins = Inserter(doc, api10)
        ins.para("Convention: Guest endpoints need no token; others need a JWT. Shop "
                 "endpoints are membership-scoped; admin endpoints require the Admin role. "
                 "Status: implemented vs designed.")
        for title, rows in parse_api_md(API):
            ins.heading(title, 3)
            ins.table(rows, sizes=[8] * len(rows[0]))

    doc.save(OUT)
    print("Wrote:", OUT)
    print("Tables:", len(doc.tables))
    print("Open in Word, Ctrl+A then F9 to refresh the Table of Contents.")


if __name__ == "__main__":
    main()
